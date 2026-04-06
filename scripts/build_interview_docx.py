"""Convert docs/PROJECT_WALKTHROUGH_INTERVIEW.md → .docx for interview prep (run from repo root)."""

from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.shared import Pt

ROOT = Path(__file__).resolve().parent.parent
MD_PATH = ROOT / "docs" / "PROJECT_WALKTHROUGH_INTERVIEW.md"
OUT_PATH = ROOT / "docs" / "PROJECT_WALKTHROUGH_INTERVIEW.docx"


def add_runs(paragraph, text: str) -> None:
    parts = re.split(r"(\*\*.+?\*\*|`[^`]+`)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            r = paragraph.add_run(part[2:-2])
            r.bold = True
        elif part.startswith("`") and part.endswith("`"):
            r = paragraph.add_run(part[1:-1])
            r.font.name = "Consolas"
            r.font.size = Pt(10)
        else:
            paragraph.add_run(part)


def parse_table_rows(lines: list[str], start: int) -> tuple[list[list[str]], int]:
    rows: list[list[str]] = []
    i = start
    while i < len(lines):
        line = lines[i].strip()
        if not line.startswith("|"):
            break
        if re.match(r"^\|\s*[-:]+\s*\|", line):
            i += 1
            continue
        cells = [c.strip() for c in line.split("|")]
        cells = [c for c in cells if c != ""]
        if cells:
            rows.append(cells)
        i += 1
    return rows, i


def main() -> None:
    if not MD_PATH.is_file():
        raise SystemExit(f"Missing {MD_PATH}")

    text = MD_PATH.read_text(encoding="utf-8")
    lines = text.splitlines()

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if not stripped:
            i += 1
            continue

        if stripped == "---":
            doc.add_paragraph()
            i += 1
            continue

        if stripped.startswith("### "):
            doc.add_heading(stripped[4:].strip(), level=2)
            i += 1
            continue

        if stripped.startswith("## "):
            doc.add_heading(stripped[3:].strip(), level=1)
            i += 1
            continue

        if stripped.startswith("# "):
            doc.add_heading(stripped[2:].strip(), level=0)
            i += 1
            continue

        if stripped.startswith("|") and "|" in stripped[1:]:
            table_rows, ni = parse_table_rows(lines, i)
            if table_rows:
                ncols = max(len(r) for r in table_rows)
                table = doc.add_table(rows=len(table_rows), cols=ncols)
                table.style = "Table Grid"
                for ri, row in enumerate(table_rows):
                    for ci in range(ncols):
                        cell_text = row[ci] if ci < len(row) else ""
                        cell = table.rows[ri].cells[ci]
                        cell.text = ""
                        p = cell.paragraphs[0]
                        add_runs(p, cell_text)
                doc.add_paragraph()
            i = ni
            continue

        if stripped.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            add_runs(p, stripped[2:].strip())
            i += 1
            continue

        if re.match(r"^\d+\.\s", stripped):
            p = doc.add_paragraph(style="List Number")
            add_runs(p, re.sub(r"^\d+\.\s*", "", stripped))
            i += 1
            continue

        p = doc.add_paragraph()
        add_runs(p, stripped)
        i += 1

    OUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    try:
        doc.save(OUT_PATH)
        print(f"Wrote {OUT_PATH}")
    except PermissionError:
        alt = OUT_PATH.with_name(f"{OUT_PATH.stem}_NEW{OUT_PATH.suffix}")
        doc.save(alt)
        print(
            f"Could not overwrite {OUT_PATH} (close Word if it is open). Wrote {alt} instead."
        )


if __name__ == "__main__":
    if str(ROOT) not in sys.path:
        sys.path.insert(0, str(ROOT))
    main()
