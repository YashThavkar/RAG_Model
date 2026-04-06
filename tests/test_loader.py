from __future__ import annotations

"""Tiny synthetic PDF/DOCX files to prove loader output shape."""

import fitz
from docx import Document

from app.loader import load_docx, load_documents, load_pdf


def test_load_pdf_page(tmp_path):
    p = tmp_path / "t.pdf"
    doc = fitz.open()
    page = doc.new_page()
    page.insert_text((72, 72), "Hello Transformer attention.")
    doc.save(p)
    doc.close()
    units = load_pdf(p)
    assert len(units) >= 1
    assert units[0]["document_type"] == "pdf"
    assert "attention" in units[0]["raw_text"].lower()


def test_load_docx_paragraphs(tmp_path):
    p = tmp_path / "t.docx"
    d = Document()
    d.add_heading("Section A", level=1)
    d.add_paragraph("EU AI high-risk obligations include documentation.")
    d.save(p)
    units = load_docx(p)
    assert any("documentation" in u["raw_text"].lower() for u in units)


def test_load_documents_both(tmp_path):
    pdf = tmp_path / "a.pdf"
    docx = tmp_path / "b.docx"
    d = fitz.open()
    d.new_page().insert_text((72, 72), "Recurrent models are slow.")
    d.save(pdf)
    d.close()
    dx = Document()
    dx.add_paragraph("Minimal risk AI has fewer obligations.")
    dx.save(docx)
    units = load_documents([pdf, docx])
    assert len(units) >= 2
